#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""WPS 文档助手 · 用 python-docx 生成 / 读取 .docx(WPS Writer 与 Word 通用格式)。

依赖:pip install python-docx
设计为既可被 Codex 当函数库调用,也可直接运行生成一份示例周报:
    python docx_tools.py 周报.docx
"""
from __future__ import annotations
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def new_document():
    """新建空白文档。返回 Document 对象。"""
    return Document()


def add_title(doc, text, subtitle=None):
    """加大标题(居中)+ 可选副标题。"""
    h = doc.add_heading(text, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(12)
    return doc


def add_section(doc, heading, body=None, level=1):
    """加一节:小标题 + 可选正文段落(body 可为字符串或字符串列表;省略则只加标题)。"""
    doc.add_heading(heading, level=level)
    if body:
        for para in ([body] if isinstance(body, str) else body):
            doc.add_paragraph(para)
    return doc


def add_table(doc, header, rows, style="Light Grid Accent 1"):
    """加表格。header=表头列表,rows=二维列表。"""
    table = doc.add_table(rows=1, cols=len(header))
    try:
        table.style = style
    except KeyError:
        pass  # 模板缺该样式时退回默认,不报错
    for i, name in enumerate(header):
        table.rows[0].cells[i].text = str(name)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def read_document(path):
    """读取 .docx 全文文本(每段一行)。"""
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)


def build_demo(path="示例-周报.docx"):
    """生成一份示例周报,验证依赖与流程是否打通。"""
    doc = new_document()
    add_title(doc, "周 工 作 报 告", subtitle="研发部 · 2026 年第 23 周")
    add_section(doc, "一、本周完成", [
        "完成 Codex 中文插件市场导航站的样式重构与上线准备。",
        "新增 WPS 办公三件套技能,并通过真机测试。",
    ])
    add_section(doc, "二、数据概览")
    add_table(
        doc,
        header=["指标", "本周", "环比"],
        rows=[["新增用户", "1,280", "+12%"], ["活跃天数", "5", "持平"], ["反馈处理", "37", "+8"]],
    )
    add_section(doc, "三、下周计划", ["推进 WPS 表格与演示技能的模板库。", "完善提交收录流程。"])
    doc.save(path)
    return path


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else "示例-周报.docx"
    saved = build_demo(out)
    print("已生成 WPS/Word 文档:", saved)
    print("--- 读回校验(前 3 行)---")
    print("\n".join(read_document(saved).splitlines()[:3]))
